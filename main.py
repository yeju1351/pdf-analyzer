import streamlit as st
import pdfplumber
import google.generativeai as genai
from docx import Document
import io
import olefile
import zlib

# 서버의 비밀 공간(secrets)에서 키를 몰래 가져옵니다.
API_KEY = st.secrets["GOOGLE_API_KEY"]
genai.configure(api_key=API_KEY)

# 1. PDF 텍스트 추출 함수
def get_text_from_pdf(file):
    all_text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                all_text += extracted + "\n"
    return all_text

# 2. 한글(HWP) 텍스트 추출 함수
def get_text_from_hwp(file):
    try:
        file_bytes = file.read()
        ole = olefile.OleFileIO(file_bytes)
        dirs = ole.listdir()
        valid_dirs = [d for d in dirs if d[0] == "BodyText"]
        text = ""
        for d in valid_dirs:
            stream = ole.openstream(d)
            data = stream.read()
            try:
                decoded = zlib.decompress(data, -15)
                text += decoded.decode('utf-16le', errors='ignore')
            except Exception:
                pass
        file.seek(0)
        return text
    except Exception as e:
        return f"\n[HWP 추출 오류: {e}]\n"

# 3. 통합 파일 추출 함수
def get_text_from_file(file):
    file_name = file.name.lower()
    if file_name.endswith('.pdf'):
        return get_text_from_pdf(file)
    elif file_name.endswith('.hwp'):
        return get_text_from_hwp(file)
    else:
        return ""

# 4. 💡 AI 분석 함수 (표 양식으로 완벽 고정)
def analyze_text_with_ai(text):
    model = genai.GenerativeModel('gemini-2.5-flash')
    prompt = f"""
    다음은 사업 공고문, 세부평가기준, 작성안내서 등의 전체 문서 내용입니다. 
    이 내용을 종합적으로 분석하여, 아래의 [출력 양식]과 완벽하게 동일한 형태의 표(마크다운 Table)로 정리해 주세요.
    HWP 파일 텍스트 추출 과정에서 이상한 기호가 섞여 있더라도 무시하고 전체 문맥을 파악해 주세요.

    [분석 지시사항]
    1. 문서를 꼼꼼히 읽고 '실격, 평가제외, 감점, 최하점, 불가, 탈락, 주의사항'에 해당하는 조건들을 모두 찾아내세요.
    2. 찾은 조건들을 '구분, 조건유형, 세부 내용, 출처' 4개의 열로 분류하여 표 형식으로 작성하세요.
    3. '출처'에는 해당 내용이 어느 문서의 어느 부분(예: 공고문 p.2, 세부평가기준 Ⅳ.6)에 있는지 최대한 정확히 기재하세요.

    [출력 양식 예시]
    | 구분 | 조건유형 | 세부 내용 | 출처 |
    | :--- | :--- | :--- | :--- |
    | 공동도급 | 주의사항 | 자격1) + 자격2) 업체의 공동도급 형태로만 참여 가능(단독 참여 불가) | 공고문 p.1 |
    | 책임기술인 | 감점 | A4로 작성해야 하며, A3로 작성한 경우 2Page로 계산. 초과 시 0.5점 감점 | 작성안내서 붙임1 |

    [전체 문서 내용]
    {text}
    """
    response = model.generate_content(prompt)
    return response.text

# 5. 💡 워드 파일 생성 함수 (진짜 표 테두리 그리기 기능 추가)
def create_word_file(result_text):
    doc = Document()
    doc.add_heading('📄 AI 사업 공고 분석 보고서', 0)
    
    table = None
    for line in result_text.split('\n'):
        line = line.strip()
        
        # 마크다운 표의 구분선(|---|---|)은 워드에서 그릴 필요가 없으므로 무시합니다.
        if line.startswith('|') and line.endswith('|') and '-' in line:
            if line.replace('|', '').replace('-', '').replace(' ', '').replace(':', '') == '':
                continue
        
        # 표 데이터 추출 및 워드 표 생성
        if line.startswith('|') and line.endswith('|'):
            row_data = [cell.strip() for cell in line.strip('|').split('|')]
            
            if table is None:
                table = doc.add_table(rows=1, cols=len(row_data))
                table.style = 'Table Grid' # 워드 기본 표 테두리 스타일
                row_cells = table.rows[0].cells
                for i in range(len(row_data)):
                    row_cells[i].text = row_data[i]
            else:
                row_cells = table.add_row().cells
                for i in range(min(len(row_data), len(row_cells))):
                    row_cells[i].text = row_data[i]
        else:
            table = None 
            if line:
                doc.add_paragraph(line)
                
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =====================================================================
# --- 웹 화면(UI) 구성 ---
# =====================================================================

st.set_page_config(page_title="AI 공고문 표 분석기", page_icon="📊")

st.title("📊 AI 사업 공고문 표 분석기")

# 💡 동료들을 위한 무료 API 안내 문구 추가
st.info("🚨 **[이용 안내]** 무료 AI 엔진을 사용 중이므로 한 번에 많은 요청이 몰리면 에러가 뜰 수 있습니다. 에러 발생 시 **1~2분 뒤 새로고침(F5)**하여 다시 시도해 주세요!")

st.write("관련 PDF와 한글(HWP) 파일들을 한 번에 올리면, AI가 핵심 조건을 '표'로 완벽하게 정리해 드립니다!")

uploaded_files = st.file_uploader("📂 여기에 파일들을 마우스로 끌어다 놓으세요 (PDF, HWP)", type=["pdf", "hwp"], accept_multiple_files=True)

if uploaded_files:
    if st.button("✨ 표 분석 시작하기"):
        
        combined_text = ""
        with st.spinner('AI가 열심히 문서를 읽고 표를 그리고 있습니다. 잠시만 기다려주세요... ☕'):
            
            for file in uploaded_files:
                combined_text += f"\n\n--- [{file.name} 문서 내용] ---\n\n"
                combined_text += get_text_from_file(file)
            
            result = analyze_text_with_ai(combined_text)
            
            st.success("🎉 분석이 완료되었습니다! 결과를 확인해 주세요.")
            st.markdown(result)
            
            st.divider()
            st.subheader("💾 보고서 파일로 다운로드")
            col1, col2 = st.columns(2)
            
            with col1:
                st.download_button(
                    label="📝 텍스트(.txt)로 저장",
                    data=result,
                    file_name="공고문_표분석.txt",
                    mime="text/plain"
                )
            with col2:
                # 💡 새로 만든 워드 표 그리기 함수를 연결했습니다.
                word_data = create_word_file(result)
                st.download_button(
                    label="📄 워드(.docx)로 표 저장",
                    data=word_data,
                    file_name="공고문_표분석.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )